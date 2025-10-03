#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Document Processing MCP Server

Core Tools (5):
1. extract_docx_images(docx_path, user_working_dir, output_dir="pictures")
2. extract_docx_text(docx_path, user_working_dir, max_chars=50000) - 现在包含表格和Excel内容
3. extract_docx_tables(docx_path, user_working_dir, include_excel=True) - 专门提取表格和Excel
4. extract_zip_assets(zip_path, user_working_dir, output_dir="pictures")
5. tag_exported_images(image_dir, user_working_dir, ocr_lang="chi_sim+eng")

新增功能:
- 提取Word文档中的表格内容
- 提取嵌入在Word文档中的Excel文件内容
- 支持多工作表Excel文件
- 表格数据以结构化格式返回

IMPORTANT FOR AI USAGE:
- user_working_dir is REQUIRED for all tools
- This ensures consistent path resolution
- AI must provide the user's project directory

HOW TO GET USER WORKING DIRECTORY:
1. Ask the user: "What is your current project directory?"
2. Use the directory where the user is working (e.g., where their files are located)
3. Provide the full absolute path (e.g., "D:\\Users\\username\\project")
4. If user provides relative path, convert to absolute path
5. Always verify the directory exists before calling tools

EXAMPLE USAGE:
- User says: "Extract images from document.docx"
- AI should ask: "What is your current project directory?"
- User responds: "D:\\MyProject"
- AI calls: extract_docx_images("document.docx", "D:\\MyProject", "pictures")

- User says: "Extract tables from document.docx"
- AI calls: extract_docx_tables("document.docx", "D:\\MyProject", include_excel=True)

Requirements: 
- pip install mcp python-docx pillow pytesseract pandas openpyxl
- For OCR (optional): install Tesseract and language data (e.g., chi_sim)
- For Excel support: pandas and openpyxl are required
"""

import io
import os
import zipfile
from pathlib import Path

from PIL import Image
from mcp.server.fastmcp import FastMCP

try:
    from docx import Document
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False

try:
    import pytesseract
    OCR_AVAILABLE = True
except Exception:
    OCR_AVAILABLE = False

try:
    import pandas as pd
    import openpyxl
    TABLES_AVAILABLE = True
except Exception:
    TABLES_AVAILABLE = False

# Initialize FastMCP server
mcp = FastMCP("DocxImageTagger")


def _ensure_dir(p: Path):
    p.mkdir(parents=True, exist_ok=True)


def _bytes_to_image_info(b: bytes):
    try:
        with Image.open(io.BytesIO(b)) as im:
            return {"format": im.format, "width": im.width, "height": im.height}
    except Exception:
        return {"format": "unknown", "width": None, "height": None}


def _ocr_image(path: Path, lang: str = "eng"):
    if not OCR_AVAILABLE:
        return ""
    try:
        with Image.open(path) as img:
            return pytesseract.image_to_string(img, lang=lang).strip()
    except Exception:
        return ""


def _extract_docx_tables(docx_path: Path) -> list:
    """从docx文件中提取表格内容"""
    if not DOCX_AVAILABLE:
        return []
    
    try:
        doc = Document(str(docx_path))
        tables = []
        
        for table_idx, table in enumerate(doc.tables):
            table_data = []
            for row_idx, row in enumerate(table.rows):
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    row_data.append(cell_text)
                table_data.append(row_data)
            
            if table_data:  # 只添加非空表格
                tables.append({
                    "table_index": table_idx + 1,
                    "rows": len(table_data),
                    "columns": len(table_data[0]) if table_data else 0,
                    "data": table_data
                })
        
        return tables
    except Exception as e:
        return [{"error": f"Failed to extract tables: {str(e)}"}]


def _extract_excel_from_docx(docx_path: Path) -> list:
    """从docx文件中提取嵌入的Excel文件"""
    if not TABLES_AVAILABLE:
        return []
    
    try:
        excel_files = []
        with zipfile.ZipFile(docx_path, "r") as zf:
            # 查找嵌入的Excel文件
            excel_names = [n for n in zf.namelist() 
                          if n.lower().endswith(('.xlsx', '.xls')) and 
                          'word/embeddings/' in n.lower()]
            
            for excel_name in excel_names:
                try:
                    excel_data = zf.read(excel_name)
                    with zipfile.ZipFile(io.BytesIO(excel_data), "r") as excel_zip:
                        # 读取Excel文件内容
                        workbook = openpyxl.load_workbook(io.BytesIO(excel_data))
                        sheets_data = []
                        
                        for sheet_name in workbook.sheetnames:
                            sheet = workbook[sheet_name]
                            sheet_data = []
                            
                            for row in sheet.iter_rows(values_only=True):
                                if any(cell is not None for cell in row):  # 跳过空行
                                    sheet_data.append([str(cell) if cell is not None else "" for cell in row])
                            
                            if sheet_data:
                                sheets_data.append({
                                    "sheet_name": sheet_name,
                                    "rows": len(sheet_data),
                                    "columns": len(sheet_data[0]) if sheet_data else 0,
                                    "data": sheet_data
                                })
                        
                        excel_files.append({
                            "filename": excel_name,
                            "sheets": sheets_data
                        })
                        
                except Exception as e:
                    excel_files.append({
                        "filename": excel_name,
                        "error": f"Failed to read Excel file: {str(e)}"
                    })
        
        return excel_files
    except Exception as e:
        return [{"error": f"Failed to extract Excel files: {str(e)}"}]


def _validate_path(path_str: str, expected_ext: str = None) -> dict:
    """验证路径并返回标准化结果"""
    try:
        path = Path(path_str).resolve()
        if not path.exists():
            return {"valid": False, "error": f"Path does not exist: {path_str}"}
        if expected_ext and path.suffix.lower() != expected_ext.lower():
            return {"valid": False, "error": f"Expected {expected_ext} file, got {path.suffix}"}
        return {"valid": True, "path": path, "absolute_path": str(path)}
    except Exception as e:
        return {"valid": False, "error": f"Invalid path: {str(e)}"}


@mcp.tool()
def extract_docx_images(docx_path: str, user_working_dir: str, output_dir: str = "pictures") -> dict:
    """Extract images from a .docx file into output_dir. Return list and meta.
    
    IMPORTANT: AI must ask user for their project directory before calling this tool.
    
    Args:
        docx_path: Path to the .docx file (relative to user_working_dir if not absolute)
        user_working_dir: User's working directory (REQUIRED - ask user for this)
        output_dir: Output directory for extracted images (default: "pictures")
    
    Returns:
        dict: Contains count, images list, and success status
    """
    try:
        # 确定基础目录
        base_dir = Path(user_working_dir)
        if not base_dir.exists():
            return {"error": f"User working directory not found: {user_working_dir}"}
        
        # 标准化输入路径
        if Path(docx_path).is_absolute():
            dp = Path(docx_path)
        else:
            dp = base_dir / docx_path
        
        if not dp.exists():
            return {
                "error": f"File not found: {docx_path}", 
                "suggestion": f"Tried: {docx_path} in {base_dir}",
                "base_directory": str(base_dir)
            }
        if dp.suffix.lower() != ".docx":
            return {"error": f"Not a .docx file: {docx_path}", "suggestion": "Please provide a valid .docx file"}
        
        # 标准化输出路径
        if Path(output_dir).is_absolute():
            output_path = Path(output_dir)
        else:
            output_path = base_dir / output_dir
        _ensure_dir(output_path)
        
        saved = []
        with zipfile.ZipFile(dp, "r") as zf:
            names = [n for n in zf.namelist() if n.startswith("word/media/")]
            if not names:
                return {"count": 0, "images": [], "message": "No images found in the document"}
            
            for i, name in enumerate(names, 1):
                data = zf.read(name)
                info = _bytes_to_image_info(data)
                ext = os.path.splitext(name)[1] or ".png"
                out_name = f"docx_img_{i:03d}{ext}"
                out_path = output_path / out_name
                
                with open(out_path, "wb") as f:
                    f.write(data)
                saved.append({
                    "filename": str(out_path),
                    "relative_path": str(out_path.relative_to(output_path)),
                    **info
                })
        
        return {
            "count": len(saved),
            "images": saved,
            "output_directory": str(output_path),
            "success": True,
            "message": f"Successfully extracted {len(saved)} images to {output_path}"
        }
        
    except zipfile.BadZipFile:
        return {"error": "Invalid .docx file format", "suggestion": "The file may be corrupted or not a valid Word document"}
    except PermissionError:
        return {"error": "Permission denied", "suggestion": "Please check if you have write permissions to the output directory"}
    except Exception as e:
        return {"error": f"Unexpected error: {str(e)}", "suggestion": "Please try again or contact support"}


@mcp.tool()
def extract_docx_text(docx_path: str, user_working_dir: str, max_chars: int = 50000) -> dict:
    """Extract text from .docx (preview limited by max_chars).
    
    IMPORTANT: AI must ask user for their project directory before calling this tool.
    
    Args:
        docx_path: Path to the .docx file (relative to user_working_dir if not absolute)
        user_working_dir: User's working directory (REQUIRED - ask user for this)
        max_chars: Maximum characters to extract (default: 50000)
    
    Returns:
        dict: Contains text content, character count, and success status
    """
    try:
        # 确定基础目录
        base_dir = Path(user_working_dir)
        if not base_dir.exists():
            return {"error": f"User working directory not found: {user_working_dir}"}
        
        # 标准化输入路径
        if Path(docx_path).is_absolute():
            dp = Path(docx_path)
        else:
            dp = base_dir / docx_path
        
        if not dp.exists():
            return {
                "error": f"File not found: {docx_path}", 
                "suggestion": f"Tried: {docx_path} in {base_dir}",
                "base_directory": str(base_dir)
            }
        if dp.suffix.lower() != ".docx":
            return {"error": f"Not a .docx file: {docx_path}", "suggestion": "Please provide a valid .docx file"}
        if not DOCX_AVAILABLE:
            return {"error": "python-docx not installed", "suggestion": "Please install python-docx: pip install python-docx"}

        doc = Document(str(dp))
        text = "\n".join(p.text for p in doc.paragraphs if p.text)
        
        # 提取表格内容
        tables = _extract_docx_tables(dp)
        
        # 提取嵌入的Excel文件
        excel_files = _extract_excel_from_docx(dp)
        
        # 构建完整内容
        full_content = text
        content_parts = []
        
        if text.strip():
            content_parts.append(f"=== 文本内容 ===\n{text}")
        
        if tables:
            content_parts.append("=== 表格内容 ===")
            for table in tables:
                if "error" not in table:
                    content_parts.append(f"\n表格 {table['table_index']} ({table['rows']}行 x {table['columns']}列):")
                    for row in table['data']:
                        content_parts.append(" | ".join(row))
                else:
                    content_parts.append(f"表格提取错误: {table['error']}")
        
        if excel_files:
            content_parts.append("=== 嵌入的Excel文件 ===")
            for excel_file in excel_files:
                if "error" not in excel_file:
                    content_parts.append(f"\nExcel文件: {excel_file['filename']}")
                    for sheet in excel_file['sheets']:
                        content_parts.append(f"  工作表: {sheet['sheet_name']} ({sheet['rows']}行 x {sheet['columns']}列)")
                        for row in sheet['data'][:5]:  # 只显示前5行
                            content_parts.append("    " + " | ".join(row))
                        if len(sheet['data']) > 5:
                            content_parts.append(f"    ... (还有{len(sheet['data'])-5}行)")
                else:
                    content_parts.append(f"Excel文件读取错误: {excel_file['error']}")
        
        full_content = "\n".join(content_parts)
        
        if not full_content.strip():
            return {"chars": 0, "preview": "", "message": "No content found in the document"}
        
        return {
            "chars": len(full_content),
            "preview": full_content[:max_chars],
            "success": True,
            "message": f"Successfully extracted {len(full_content)} characters",
            "truncated": len(full_content) > max_chars,
            "tables_count": len(tables),
            "excel_files_count": len(excel_files),
            "tables": tables,
            "excel_files": excel_files
        }
        
    except Exception as e:
        return {"error": f"Failed to extract text: {str(e)}", "suggestion": "Please check if the file is a valid Word document"}


@mcp.tool()
def extract_docx_tables(docx_path: str, user_working_dir: str, include_excel: bool = True) -> dict:
    """专门提取docx文件中的表格和Excel内容
    
    IMPORTANT: AI must ask user for their project directory before calling this tool.
    
    Args:
        docx_path: Path to the .docx file (relative to user_working_dir if not absolute)
        user_working_dir: User's working directory (REQUIRED - ask user for this)
        include_excel: Whether to extract embedded Excel files (default: True)
    
    Returns:
        dict: Contains tables and Excel files data
    """
    try:
        # 确定基础目录
        base_dir = Path(user_working_dir)
        if not base_dir.exists():
            return {"error": f"User working directory not found: {user_working_dir}"}
        
        # 标准化输入路径
        if Path(docx_path).is_absolute():
            dp = Path(docx_path)
        else:
            dp = base_dir / docx_path
        
        if not dp.exists():
            return {
                "error": f"File not found: {docx_path}", 
                "suggestion": f"Tried: {docx_path} in {base_dir}",
                "base_directory": str(base_dir)
            }
        if dp.suffix.lower() != ".docx":
            return {"error": f"Not a .docx file: {docx_path}", "suggestion": "Please provide a valid .docx file"}
        if not DOCX_AVAILABLE:
            return {"error": "python-docx not installed", "suggestion": "Please install python-docx: pip install python-docx"}

        # 提取Word表格
        tables = _extract_docx_tables(dp)
        
        # 提取嵌入的Excel文件
        excel_files = []
        if include_excel:
            excel_files = _extract_excel_from_docx(dp)
        
        return {
            "success": True,
            "message": f"Successfully extracted {len(tables)} tables and {len(excel_files)} Excel files",
            "tables_count": len(tables),
            "excel_files_count": len(excel_files),
            "tables": tables,
            "excel_files": excel_files,
            "tables_available": DOCX_AVAILABLE,
            "excel_available": TABLES_AVAILABLE
        }
        
    except Exception as e:
        return {"error": f"Failed to extract tables: {str(e)}", "suggestion": "Please check if the file is a valid Word document"}


@mcp.tool()
def extract_zip_assets(zip_path: str, user_working_dir: str, output_dir: str = "pictures") -> dict:
    """
    Extract images from a .zip file. If the zip contains .docx files, also
    extract images from each docx found. Images are placed into output_dir.
    
    IMPORTANT: AI must ask user for their project directory before calling this tool.
    
    Args:
        zip_path: Path to the .zip file (relative to user_working_dir if not absolute)
        user_working_dir: User's working directory (REQUIRED - ask user for this)
        output_dir: Output directory for extracted images
    """
    try:
        # 确定基础目录
        base_dir = Path(user_working_dir)
        if not base_dir.exists():
            return {"error": f"User working directory not found: {user_working_dir}"}
        
        # 标准化输入路径
        if Path(zip_path).is_absolute():
            zp = Path(zip_path)
        else:
            zp = base_dir / zip_path
        
        if not zp.exists():
            return {"error": f"File not found: {zip_path}", "suggestion": f"Tried: {zip_path} in {base_dir}"}
        if zp.suffix.lower() != ".zip":
            return {"error": f"Not a .zip file: {zip_path}", "suggestion": "Please provide a valid .zip file"}
        
        # 标准化输出路径
        if Path(output_dir).is_absolute():
            out = Path(output_dir)
        else:
            out = base_dir / output_dir
        _ensure_dir(out)

        total = 0
        details = []
        with zipfile.ZipFile(zp, "r") as zf:
            names = zf.namelist()
            # Direct image files in zip
            for n in names:
                if n.lower().endswith((".png", ".jpg", ".jpeg", ".bmp", ".gif", ".tif", ".tiff", ".webp")):
                    data = zf.read(n)
                    info = _bytes_to_image_info(data)
                    ext = os.path.splitext(n)[1]
                    out_name = f"zip_img_{total+1:03d}{ext}"
                    out_path = out / out_name
                    with open(out_path, "wb") as f:
                        f.write(data)
                    total += 1
                    details.append({"from": n, "filename": str(out_path), **info})

            # Embedded .docx files in zip
            for n in names:
                if n.lower().endswith(".docx"):
                    try:
                        b = zf.read(n)
                        with zipfile.ZipFile(io.BytesIO(b), "r") as dzip:
                            media = [m for m in dzip.namelist() if m.startswith("word/media/")]
                            for m in media:
                                data = dzip.read(m)
                                info = _bytes_to_image_info(data)
                                ext = os.path.splitext(m)[1] or ".png"
                                out_name = f"zip_docx_img_{total+1:03d}{ext}"
                                out_path = out / out_name
                                with open(out_path, "wb") as f:
                                    f.write(data)
                                total += 1
                                details.append({"from_docx": n, "entry": m, "filename": str(out_path), **info})
                    except Exception as e:
                        details.append({"docx_in_zip": n, "error": str(e)})

        return {
            "count": total, 
            "images": details,
            "output_directory": str(out),
            "success": True,
            "message": f"Successfully extracted {total} images to {out}"
        }
        
    except Exception as e:
        return {"error": f"Failed to extract zip assets: {str(e)}", "suggestion": "Please check if the file is a valid zip archive"}


@mcp.tool()
def tag_exported_images(image_dir: str, user_working_dir: str, ocr_lang: str = "chi_sim+eng") -> dict:
    """Assign simple tags (format/size + OCR preview if available) to images in a directory.
    
    IMPORTANT: AI must ask user for their project directory before calling this tool.
    
    Args:
        image_dir: Directory containing images to tag (relative to user_working_dir if not absolute)
        user_working_dir: User's working directory (REQUIRED - ask user for this)
        ocr_lang: OCR language code (default: "chi_sim+eng")
    """
    try:
        # 确定基础目录
        base_dir = Path(user_working_dir)
        if not base_dir.exists():
            return {"error": f"User working directory not found: {user_working_dir}"}
        
        # 标准化路径
        if Path(image_dir).is_absolute():
            p = Path(image_dir)
        else:
            p = base_dir / image_dir
        
        if not p.exists():
            return {"error": f"Directory not found: {image_dir}", "suggestion": f"Tried: {image_dir} in {base_dir}"}
        if not p.is_dir():
            return {"error": f"Not a directory: {image_dir}", "suggestion": "Please provide a directory path instead of a file"}
        
        items = []
        for fn in sorted(p.iterdir()):
            if not fn.is_file():
                continue
            if fn.suffix.lower() not in [".png", ".jpg", ".jpeg", ".bmp", ".gif", ".tif", ".tiff", ".webp"]:
                continue
            try:
                with open(fn, "rb") as f:
                    b = f.read()
                info = _bytes_to_image_info(b)
                ocr_text = _ocr_image(fn, lang=ocr_lang) if OCR_AVAILABLE else ""
                tags = []
                if info["format"]:
                    tags.append(info["format"].lower())
                if info["width"] and info["height"]:
                    tags.append(f"{info['width']}x{info['height']}")
                if ocr_text:
                    tags.append("ocr")
                items.append({
                    "file": str(fn),
                    "format": info["format"],
                    "width": info["width"],
                    "height": info["height"],
                    "ocr_preview": (ocr_text[:160] + "…") if len(ocr_text) > 160 else ocr_text,
                    "tags": tags
                })
            except Exception as e:
                items.append({"file": str(fn), "error": str(e)})
        
        return {
            "count": len(items), 
            "items": items, 
            "ocr_available": OCR_AVAILABLE,
            "directory": str(p),
            "success": True,
            "message": f"Successfully processed {len(items)} images in {p}"
        }
        
    except Exception as e:
        return {"error": f"Failed to process images: {str(e)}", "suggestion": "Please check the directory path and permissions"}


if __name__ == "__main__":
    mcp.run(transport="stdio")